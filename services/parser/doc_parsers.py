import io
from typing import Tuple, List, Dict, Any
from pypdf import PdfReader
from docx import Document as DocxDocument
from core.models import Document, DocumentChunk
from services.parser.base import BaseParser
from ingestion.chunkers import SemanticAPIChunker
from core.exceptions import ParsingException


class PDFParser(BaseParser):
    """
    Parser for PDF documents.
    Extracts text page-by-page to preserve page number metadata during chunking.
    """

    def __init__(self, chunker: SemanticAPIChunker = None):
        # Default chunker configured with global settings
        self.chunker = chunker or SemanticAPIChunker()

    def parse(self, file_content: bytes, filename: str) -> Tuple[Document, List[DocumentChunk]]:
        try:
            # Create parent document model
            doc = Document(filename=filename, file_type="pdf")
            chunks: List[DocumentChunk] = []

            # Load PDF from bytes stream
            pdf_stream = io.BytesIO(file_content)
            reader = PdfReader(pdf_stream)
            num_pages = len(reader.pages)
            
            doc.metadata["total_pages"] = num_pages

            # Iterate page-by-page to preserve page numbers
            for page_idx in range(num_pages):
                page = reader.pages[page_idx]
                page_text = page.extract_text()
                
                if not page_text or not page_text.strip():
                    continue

                # Normalize text
                normalized_text = self._normalize(page_text)
                if not normalized_text:
                    continue

                # Metadata specific to this page's chunks
                page_metadata = {
                    "filename": filename,
                    "page_number": page_idx + 1,
                    "total_pages": num_pages
                }

                # Chunk the page text
                page_chunks = self.chunker.chunk_text(
                    text=normalized_text,
                    document_id=doc.id,
                    base_metadata=page_metadata
                )
                
                # Update chunk index across the document
                for chunk in page_chunks:
                    chunk.chunk_index = len(chunks)
                    chunks.append(chunk)

            return doc, chunks

        except Exception as e:
            self._handle_error(filename, e)


class DocxParser(BaseParser):
    """
    Parser for Microsoft Word (.docx) documents.
    Tracks headings to append section names to chunk metadata.
    """

    def __init__(self, chunker: SemanticAPIChunker = None):
        self.chunker = chunker or SemanticAPIChunker()

    def parse(self, file_content: bytes, filename: str) -> Tuple[Document, List[DocumentChunk]]:
        try:
            doc = Document(filename=filename, file_type="docx")
            chunks: List[DocumentChunk] = []

            # Load Word doc from bytes
            docx_stream = io.BytesIO(file_content)
            word_doc = DocxDocument(docx_stream)

            current_section = "Introduction"
            section_paragraphs = []

            def process_text_block(text_block: str, section_title: str) -> List[DocumentChunk]:
                normalized = self._normalize(text_block)
                if not normalized:
                    return []
                block_metadata = {
                    "filename": filename,
                    "section": section_title
                }
                return self.chunker.chunk_text(
                    text=normalized,
                    document_id=doc.id,
                    base_metadata=block_metadata
                )

            # Iterate through paragraphs to track headers
            for para in word_doc.paragraphs:
                text = para.text.strip()
                if not text:
                    continue

                # Detect headings
                is_heading = para.style.name.startswith("Heading") or para.style.name in ["Title", "Subtitle"]
                
                if is_heading:
                    # If we already have accumulated paragraphs in a section, chunk them first
                    if section_paragraphs:
                        combined_text = "\n".join(section_paragraphs)
                        section_chunks = process_text_block(combined_text, current_section)
                        for chunk in section_chunks:
                            chunk.chunk_index = len(chunks)
                            chunks.append(chunk)
                        section_paragraphs = []
                    
                    # Update current section title
                    current_section = text
                else:
                    section_paragraphs.append(text)

            # Process remaining accumulated text
            if section_paragraphs:
                combined_text = "\n".join(section_paragraphs)
                section_chunks = process_text_block(combined_text, current_section)
                for chunk in section_chunks:
                    chunk.chunk_index = len(chunks)
                    chunks.append(chunk)

            # Process tables (tables contain highly valuable technical specification details)
            for table_idx, table in enumerate(word_doc.tables):
                table_lines = []
                for row in table.rows:
                    row_cells = [cell.text.strip() for cell in row.cells]
                    table_lines.append(" | ".join(row_cells))
                
                if table_lines:
                    table_text = f"\n[Table {table_idx + 1}]\n" + "\n".join(table_lines)
                    table_chunks = process_text_block(table_text, f"Table Data (Table {table_idx + 1})")
                    for chunk in table_chunks:
                        chunk.chunk_index = len(chunks)
                        chunks.append(chunk)

            return doc, chunks

        except Exception as e:
            self._handle_error(filename, e)


class TxtParser(BaseParser):
    """
    Parser for plain text (.txt) files.
    """

    def __init__(self, chunker: SemanticAPIChunker = None):
        self.chunker = chunker or SemanticAPIChunker()

    def parse(self, file_content: bytes, filename: str) -> Tuple[Document, List[DocumentChunk]]:
        try:
            doc = Document(filename=filename, file_type="txt")
            
            # Decode content
            raw_text = file_content.decode("utf-8", errors="ignore")
            normalized_text = self._normalize(raw_text)

            base_metadata = {"filename": filename}
            
            chunks = self.chunker.chunk_text(
                text=normalized_text,
                document_id=doc.id,
                base_metadata=base_metadata
            )
            
            return doc, chunks

        except Exception as e:
            self._handle_error(filename, e)
